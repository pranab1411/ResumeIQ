"""
modules/resume_builder.py
Unified Resume Builder Engine for ResumeIQ v2.0.
Supports 11 curated templates (ATS Classic, Modern, Corporate, Executive, Minimal,
Academic, Creative, Fresher, Experienced, One-page, Two-page).
Exports to DOCX, HTML, and PDF.
"""

import os
from typing import Dict, Any, List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from utils.logger import logger
from modules.template_engine import template_registry

def hex_to_rgb(hex_str: str) -> RGBColor:
    hex_str = hex_str.lstrip('#')
    if len(hex_str) == 6:
        r = int(hex_str[0:2], 16)
        g = int(hex_str[2:4], 16)
        b = int(hex_str[4:6], 16)
        return RGBColor(r, g, b)
    return RGBColor(79, 70, 229)

class ResumeBuilder:
    TEMPLATES = {
        1: "ATS Classic",
        2: "Modern",
        3: "Corporate",
        4: "Executive",
        5: "Minimal",
        6: "Academic",
        7: "Creative",
        8: "Fresher",
        9: "Experienced",
        10: "One-page",
        11: "Two-page"
    }

    @staticmethod
    def generate_html(data: Dict[str, Any], template_id: int = 1, options: Dict = None) -> str:
        """Generates HTML resume from Jinja2 template."""
        options = options or {}
        try:
            from jinja2 import Environment, FileSystemLoader
            template_dir = os.path.join(os.path.dirname(__file__), "..", "assets", "templates")
            env = Environment(loader=FileSystemLoader(template_dir))
            tpl = env.get_template("base_resume.html")
            
            t_data = template_registry.get_template_by_id(template_id) or {}
            
            context = dict(data)
            context["primary_color"] = options.get("primary_color", t_data.get("primary_color", "#4F46E5"))
            context["font_family"] = options.get("font_family", t_data.get("font_family", "Arial, sans-serif"))
            context["margin_size"] = options.get("margin_size", "0.75in")
            context["header_align"] = options.get("header_align", t_data.get("header_align", "left"))
            
            return tpl.render(**context)
        except Exception as e:
            logger.error(f"[ResumeBuilder] HTML generation error: {e}")
            raise Exception(f"Failed to generate HTML resume: {e}")

    @staticmethod
    def generate_pdf(data: Dict[str, Any], output_path: str, template_id: int = 1, options: Dict = None) -> str:
        """Generates PDF resume via HTML -> WeasyPrint or ReportLab PDF rendering."""
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        html_content = ResumeBuilder.generate_html(data, template_id, options)
        
        try:
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(output_path)
            logger.info(f"[ResumeBuilder] Generated PDF resume via WeasyPrint: {output_path}")
            return output_path
        except Exception as e:
            logger.warning(f"[ResumeBuilder] WeasyPrint unavailable ({e}), using ReportLab PDF export")
            # PDFReportGenerator fallback
            from modules.report_generator import PDFReportGenerator
            return PDFReportGenerator.generate(
                output_path=output_path,
                candidate_name=data.get("candidate_name", "Candidate"),
                filename="Resume.pdf",
                job_title=data.get("job_title", "Software Engineer"),
                ats_score=data.get("ats_score", 85.0),
                score_category="Excellent",
                matched_skills=data.get("skills", []),
                missing_skills=[],
                suggestions=[]
            )

    @staticmethod
    def generate_docx(data: Dict[str, Any], output_path: str, template_id: int = 1) -> str:
        """Generates a clean, professional DOCX resume."""
        return DocxResumeGenerator.generate_docx(data, output_path, template_id)

class DocxResumeGenerator:
    @staticmethod
    def generate_docx(data: Dict[str, Any], output_path: str, template_id: int = 1) -> str:
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc = Document()

        template = template_registry.get_template_by_id(template_id) or template_registry.get_template_by_id(1)
        font_name = template.get("font_family", "Calibri")
        primary_hex = template.get("primary_color", "#4F46E5")
        
        PRIMARY_COLOR = hex_to_rgb(primary_hex)
        DARK_TEXT = RGBColor(30, 41, 59)
        MUTED_TEXT = RGBColor(100, 116, 139)

        header_align = WD_ALIGN_PARAGRAPH.CENTER if template.get("header_align") == "center" else WD_ALIGN_PARAGRAPH.LEFT

        for section in doc.sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Header
        name = data.get("candidate_name", "CANDIDATE NAME").upper()
        p_name = doc.add_paragraph()
        p_name.alignment = header_align
        run_name = p_name.add_run(name)
        run_name.font.name = font_name
        run_name.font.size = Pt(22)
        run_name.font.bold = True
        run_name.font.color.rgb = PRIMARY_COLOR

        # Contact info
        email = data.get("email", "")
        phone = data.get("phone", "")
        linkedin = data.get("linkedin", "")
        github = data.get("github", "")

        contact_parts = [p for p in [phone, email, linkedin, github] if p and p != "Not Found"]
        contact_str = "  |  ".join(contact_parts) if contact_parts else "Email | Phone | LinkedIn | GitHub"

        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_contact = p_contact.add_run(contact_str)
        run_contact.font.name = font_name
        run_contact.font.size = Pt(9.5)
        run_contact.font.color.rgb = MUTED_TEXT
        p_contact.paragraph_format.space_after = Pt(14)

        def add_heading(text: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text.upper())
            run.font.name = font_name
            run.font.size = Pt(13)
            run.font.bold = True
            run.font.color.rgb = PRIMARY_COLOR
            return p

        def add_bullet(bold_prefix: str, text: str):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.line_spacing = 1.15
            
            if bold_prefix:
                run_b = p.add_run(bold_prefix + " ")
                run_b.font.name = font_name
                run_b.font.size = Pt(10.5)
                run_b.font.bold = True
                run_b.font.color.rgb = DARK_TEXT

            run_t = p.add_run(text)
            run_t.font.name = font_name
            run_t.font.size = Pt(10.5)
            run_t.font.color.rgb = DARK_TEXT
            return p

        # Summary
        summary = data.get("summary") or data.get("objective", "")
        if summary:
            add_heading("PROFESSIONAL SUMMARY")
            p_sum = doc.add_paragraph()
            p_sum.paragraph_format.space_after = Pt(8)
            run_sum = p_sum.add_run(summary)
            run_sum.font.name = font_name
            run_sum.font.size = Pt(10.5)
            run_sum.font.color.rgb = DARK_TEXT

        # Skills
        skills = data.get("skills", [])
        if skills:
            add_heading("TECHNICAL SKILLS")
            if isinstance(skills, list):
                add_bullet("Core Competencies:", ", ".join(skills))
            elif isinstance(skills, dict):
                for cat, slist in skills.items():
                    add_bullet(f"{cat}:", ", ".join(slist) if isinstance(slist, list) else str(slist))

        # Experience
        experience = data.get("experience", [])
        if experience:
            add_heading("WORK EXPERIENCE")
            for item in experience:
                title = item.get("title", item.get("content", "Role"))
                p_exp = doc.add_paragraph()
                run_t = p_exp.add_run(title)
                run_t.font.bold = True
                for bullet in item.get("bullets", []):
                    add_bullet("", bullet)

        # Education
        education = data.get("education", [])
        if education:
            add_heading("EDUCATION")
            for edu in education:
                degree = edu.get("degree", edu.get("content", "Degree"))
                p_edu = doc.add_paragraph()
                p_edu.add_run(degree).font.bold = True

        doc.save(output_path)
        logger.info(f"Generated DOCX resume at: {output_path}")
        return output_path
