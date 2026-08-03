import os
import pdfplumber
import docx
from utils.logger import logger

class DocumentParser:
    @staticmethod
    def extract_text(file_path: str) -> str:
        """Extracts plain text from PDF or DOCX file."""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found at path: {file_path}")
            
        ext = os.path.splitext(file_path)[1].lower()
        if ext == ".pdf":
            return DocumentParser._extract_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return DocumentParser._extract_from_docx(file_path)
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file format: {ext}. Only PDF, DOCX, and TXT are supported.")

    @staticmethod
    def _extract_from_pdf(file_path: str) -> str:
        text_content = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            full_text = "\n".join(text_content).strip()
            if not full_text:
                logger.warning(f"No text extracted from PDF: {file_path}. It might be scanned or image-based.")
            return full_text
        except Exception as e:
            logger.error(f"Error parsing PDF file {file_path}: {e}")
            raise Exception(f"Failed to read PDF file: {str(e)}")

    @staticmethod
    def _extract_from_docx(file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        paragraphs.append(" | ".join(row_text))
            return "\n".join(paragraphs).strip()
        except Exception as e:
            logger.error(f"Error parsing DOCX file {file_path}: {e}")
            raise Exception(f"Failed to read DOCX file: {str(e)}")
