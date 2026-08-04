"""
modules/parsers/docx_parser.py
DOCX parser for text and embedded table contents.
"""

import docx
from utils.logger import logger

def parse_docx(file_path: str) -> str:
    """Extracts text and table rows from DOCX files."""
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        
        # Extract text from tables as well
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    paragraphs.append(" | ".join(row_text))
                    
        return "\n".join(paragraphs).strip()
    except Exception as e:
        logger.error(f"[DOCXParser] Error parsing DOCX {file_path}: {e}")
        raise Exception(f"Failed to read DOCX file: {e}")
